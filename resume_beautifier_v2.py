"""
Resume Beautifier v2 - Simplified professional formatting
Focus:
- Clean bullet points
- Tighter spacing
- Professional look
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_section_line(paragraph):
    """Add horizontal line below section header"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '3')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def beautify_resume_v2(docx_path, output_path=None):
    """Apply clean, professional formatting"""
    doc = Document(docx_path)
    
    if output_path is None:
        output_path = docx_path
    
    # Section keywords
    section_keywords = [
        'PROFESSIONAL SUMMARY', 'SUMMARY', 'OBJECTIVE',
        'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT',
        'SKILLS', 'CORE SKILLS', 'TECHNICAL SKILLS',
        'EDUCATION', 'CERTIFICATIONS', 'PROJECTS',
        'ADDITIONAL INFORMATION', 'OTHER', 'COMMUNITY'
    ]
    
    # Track if we're in a bullet section
    in_bullet_section = False
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        text_upper = text.upper()
        
        # Skip empty
        if not text:
            paragraph.paragraph_format.space_after = Pt(2)
            continue
        
        # Check if section header
        is_section = any(kw in text_upper for kw in section_keywords)
        
        # Set font for all runs
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        if is_section:
            # Section header styling
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.bold = True
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(4)
            add_section_line(paragraph)
            in_bullet_section = True
        elif text.startswith('-') or text.startswith('•'):
            # Bullet point - clean it up
            clean_text = text.lstrip('-• ').strip()
            paragraph.clear()
            # Add proper bullet
            bullet_run = paragraph.add_run('• ')
            bullet_run.font.name = 'Arial'
            bullet_run.font.size = Pt(10)
            bullet_run.font.color.rgb = RGBColor(0, 0, 0)
            # Add text
            text_run = paragraph.add_run(clean_text)
            text_run.font.name = 'Arial'
            text_run.font.size = Pt(10)
            text_run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.15)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(2)
        else:
            # Regular text
            for run in paragraph.runs:
                run.font.size = Pt(10)
            
            # Check for job titles (contain title + company pattern)
            if any(x in text for x in ['Manager', 'Developer', 'Consultant', 'Owner', 'Director']):
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            
            # Company/location lines (contain | or , or dates)
            if any(x in text for x in ['|', ', Austin', ', Remote', '2021', '2022', '2023', '2024', '2025']):
                for run in paragraph.runs:
                    run.font.italic = True
                    run.font.size = Pt(9)
            
            if in_bullet_section:
                paragraph.paragraph_format.space_after = Pt(2)
    
    # Set document margins (0.5 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    doc.save(output_path)
    print(f"Beautified v2 saved: {output_path}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python resume_beautifier_v2.py <input.docx> [output.docx]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file
    
    beautify_resume_v2(input_file, output_file)
