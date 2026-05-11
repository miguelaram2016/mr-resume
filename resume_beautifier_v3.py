"""
Resume Beautifier v3 - Match professional template
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_section_line(paragraph):
    """Add horizontal line below section header"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '3')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def beautify_resume_v3(docx_path, output_path=None):
    """Match the professional template format"""
    doc = Document(docx_path)
    
    if output_path is None:
        output_path = docx_path
    
    section_keywords = [
        'PROFESSIONAL SUMMARY', 'SUMMARY',
        'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT',
        'SKILLS', 'CORE SKILLS', 'TECHNICAL SKILLS',
        'EDUCATION', 'CERTIFICATIONS', 'PROJECTS',
        'ADDITIONAL INFORMATION', 'OTHER', 'COMMUNITY'
    ]
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        text_upper = text.upper()
        
        if not text:
            continue
        
        # Check if section header
        is_section = any(kw in text_upper for kw in section_keywords)
        
        # Set uniform font
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        if is_section:
            # Make ALL CAPS and bold
            for run in paragraph.runs:
                run.text = run.text.upper()
                run.font.size = Pt(11)
                run.font.bold = True
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(4)
            add_section_line(paragraph)
            
        elif text.startswith('-') or text.startswith('•'):
            # Convert to dash bullets
            clean_text = text.lstrip('-• ').strip()
            paragraph.clear()
            bullet_run = paragraph.add_run('- ')
            bullet_run.font.name = 'Arial'
            bullet_run.font.size = Pt(10)
            bullet_run.font.color.rgb = RGBColor(0, 0, 0)
            text_run = paragraph.add_run(clean_text)
            text_run.font.name = 'Arial'
            text_run.font.size = Pt(10)
            text_run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.15)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(2)
            
        else:
            # Regular text - keep small
            for run in paragraph.runs:
                run.font.size = Pt(10)
            
            # Bold job titles
            if any(x in text for x in ['Manager', 'Developer', 'Consultant', 'Owner']):
                for run in paragraph.runs:
                    run.font.bold = True
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    
    doc.save(output_path)
    print(f"Beautified v3 saved: {output_path}")

if __name__ == "__main__":
    import sys
    input_file = sys.argv[1] if len(sys.argv) > 1 else "resume.docx"
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file
    beautify_resume_v3(input_file, output_file)
